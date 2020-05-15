#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sat Aug 10 21:32:34 2019

@author: teddy
"""

from docx import Document
from docx.shared import RGBColor
#from docx.dml.color import ColorFormat


def getText(filename):
    
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        
    hidden = '\n'.join(fullText) 
    
#    hidden_word = list(hidden.split('\n'))

    document = Document()
    run = document.add_paragraph().add_run(hidden)
    font = run.font
        

    font.color.rgb = RGBColor(0, 0, 0)

    
    run = document.add_paragraph().add_run('Warning! Keyword stuffing!')
    font = run.font
    
    font.color.rgb = RGBColor(255, 0, 0)
    
    document.save('Show_All.docx')
     
    return 'Done!'








